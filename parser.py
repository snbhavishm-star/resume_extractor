"""
Core resume parsing engine.

Text is extracted at the character/glyph level from PDF and DOCX files,
so parsing works identically regardless of the font, font size, or styling
used in the resume. What matters is the underlying text content and layout
(line breaks, section headers), not how large or which font it's rendered in.
"""

import re
import json
import calendar
from datetime import date
from pathlib import Path

import pdfplumber
import docx

from skills_db import SKILL_TO_CATEGORY, DEGREE_KEYWORDS, SECTION_HEADERS

EMAIL_RE = re.compile(r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}")
PHONE_RE = re.compile(r"(\+?\d{1,3}[\s-]?)?(\(?\d{3,5}\)?[\s-]?)?\d{3,4}[\s-]?\d{3,4}")
LINKEDIN_RE = re.compile(r"(https?://)?(www\.)?linkedin\.com/in/[A-Za-z0-9\-_/]+", re.I)
GITHUB_RE = re.compile(r"(https?://)?(www\.)?github\.com/[A-Za-z0-9\-_/]+", re.I)
YEAR_RANGE_RE = re.compile(
    r"(20\d{2}|19\d{2})\s*(-|to|–|—)\s*(20\d{2}|19\d{2}|present|current)", re.I
)
YEAR_RE = re.compile(r"(19|20)\d{2}")
CGPA_RE = re.compile(r"(?:cgpa|gpa)\s*[:\-]?\s*(\d\.\d{1,2})", re.I)
PERCENT_RE = re.compile(r"(\d{1,3}(?:\.\d{1,2})?)\s*%")

# --------------------------------------------------------------------------
# Month/year date-range detection & duration calculation
# --------------------------------------------------------------------------

MONTH_MAP = {
    "jan": 1, "january": 1, "feb": 2, "february": 2, "mar": 3, "march": 3,
    "apr": 4, "april": 4, "may": 5, "jun": 6, "june": 6, "jul": 7, "july": 7,
    "aug": 8, "august": 8, "sep": 9, "sept": 9, "september": 9, "oct": 10,
    "october": 10, "nov": 11, "november": 11, "dec": 12, "december": 12,
}
_MONTH_NAMES = "|".join(sorted(MONTH_MAP.keys(), key=len, reverse=True))

# Matches things like "March 2024", "Mar 2024", "03/2024", "3-2024"
SINGLE_DATE_RE = rf"(?:(?:{_MONTH_NAMES})\.?\s+\d{{4}}|\d{{1,2}}[/\-]\d{{4}}|\d{{4}})"
PRESENT_WORDS = r"(?:present|current|now|ongoing|till date|to date)"

# Matches "March 2024 - Present", "Jan 2023 to Dec 2024", "03/2023 – 05/2025"
MONTH_YEAR_RANGE_RE = re.compile(
    rf"({SINGLE_DATE_RE})\s*(?:-|to|–|—|until)\s*({SINGLE_DATE_RE}|{PRESENT_WORDS})",
    re.I,
)


def _parse_single_date(token: str, is_end: bool = False):
    """Parse a single date token ('March 2024', '03/2024', '2024', 'Present')
    into a Python date. For start dates, uses the 1st of the month.
    For end dates, uses the last day of the month (or today, if 'present')."""
    token = token.strip().lower()

    if re.fullmatch(PRESENT_WORDS, token):
        return date.today(), True  # (resolved_date, is_present)

    # "Month YYYY" e.g. "march 2024" / "mar. 2024"
    m = re.fullmatch(rf"({_MONTH_NAMES})\.?\s+(\d{{4}})", token)
    if m:
        month = MONTH_MAP[m.group(1)]
        year = int(m.group(2))
        if is_end:
            last_day = calendar.monthrange(year, month)[1]
            return date(year, month, last_day), False
        return date(year, month, 1), False

    # "MM/YYYY" or "MM-YYYY"
    m = re.fullmatch(r"(\d{1,2})[/\-](\d{4})", token)
    if m:
        month = int(m.group(1))
        year = int(m.group(2))
        if 1 <= month <= 12:
            if is_end:
                last_day = calendar.monthrange(year, month)[1]
                return date(year, month, last_day), False
            return date(year, month, 1), False

    # Bare "YYYY" — assume Jan for start, Dec for end
    m = re.fullmatch(r"(\d{4})", token)
    if m:
        year = int(m.group(1))
        if is_end:
            return date(year, 12, 31), False
        return date(year, 1, 1), False

    return None, False


def calculate_duration(start_token: str, end_token: str):
    """Given raw start/end date text (e.g. 'March 2024', 'Present'),
    returns a dict with years, months, total days, and a human-readable
    string. Returns None if the dates can't be confidently parsed."""
    start_date, _ = _parse_single_date(start_token, is_end=False)
    end_date, end_is_present = _parse_single_date(end_token, is_end=True)

    if not start_date or not end_date or end_date < start_date:
        return None

    total_days = (end_date - start_date).days + 1

    # Month-based breakdown (more intuitive for "X years Y months")
    months_total = (end_date.year - start_date.year) * 12 + (end_date.month - start_date.month)
    if end_date.day >= start_date.day:
        months_total += 1
    years, months = divmod(max(months_total, 0), 12)

    parts = []
    if years:
        parts.append(f"{years} year{'s' if years != 1 else ''}")
    if months:
        parts.append(f"{months} month{'s' if months != 1 else ''}")
    if not parts:
        parts.append(f"{total_days} day{'s' if total_days != 1 else ''}")

    return {
        "start_date": start_date.isoformat(),
        "end_date": "present" if end_is_present else end_date.isoformat(),
        "years": years,
        "months": months,
        "total_days": total_days,
        "human_readable": " ".join(parts),
    }


# --------------------------------------------------------------------------
# Text extraction (font-size agnostic — works on raw text, not glyph size)
# --------------------------------------------------------------------------

def extract_text_from_pdf(path: str) -> str:
    lines = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text() or ""
            lines.append(page_text)
    return "\n".join(lines)


def extract_text_from_docx(path: str) -> str:
    document = docx.Document(path)
    lines = [p.text for p in document.paragraphs]
    # also pull table content (many resumes use tables for layout)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    lines.append(cell.text)
    return "\n".join(lines)


def extract_text(path: str) -> str:
    ext = Path(path).suffix.lower()
    if ext == ".pdf":
        return extract_text_from_pdf(path)
    elif ext in (".docx", ".doc"):
        return extract_text_from_docx(path)
    elif ext == ".txt":
        return Path(path).read_text(errors="ignore")
    else:
        raise ValueError(f"Unsupported file type: {ext}")


# --------------------------------------------------------------------------
# Section segmentation
# --------------------------------------------------------------------------

def segment_sections(text: str):
    """Split resume text into named sections based on common headers."""
    lines = [l.strip() for l in text.split("\n")]
    header_positions = []

    for i, line in enumerate(lines):
        clean = re.sub(r"[^a-zA-Z& ]", "", line).strip().lower()
        if not clean or len(clean) > 40:
            continue
        for section, keywords in SECTION_HEADERS.items():
            if clean in keywords or any(clean == kw for kw in keywords):
                header_positions.append((i, section))
                break

    sections = {}
    for idx, (line_no, section) in enumerate(header_positions):
        start = line_no + 1
        end = header_positions[idx + 1][0] if idx + 1 < len(header_positions) else len(lines)
        sections[section] = "\n".join(lines[start:end]).strip()

    return sections, lines


# --------------------------------------------------------------------------
# Field extractors
# --------------------------------------------------------------------------

def extract_name(lines):
    """Heuristic: the name is usually the first substantial line without
    an email/phone/URL and without common resume keywords."""
    ignore_words = {"resume", "curriculum vitae", "cv", "profile"}
    for line in lines[:8]:
        clean = line.strip()
        if not clean:
            continue
        if EMAIL_RE.search(clean) or PHONE_RE.search(clean) or "http" in clean.lower():
            continue
        if clean.lower() in ignore_words:
            continue
        if len(clean.split()) <= 5 and len(clean) < 45:
            return clean
    return None


def extract_contact_info(text):
    email = EMAIL_RE.search(text)
    phone = PHONE_RE.search(text)
    linkedin = LINKEDIN_RE.search(text)
    github = GITHUB_RE.search(text)
    return {
        "email": email.group(0) if email else None,
        "phone": phone.group(0).strip() if phone else None,
        "linkedin": linkedin.group(0) if linkedin else None,
        "github": github.group(0) if github else None,
    }


def extract_skills(text):
    text_lower = text.lower()
    found = {}
    for skill, category in SKILL_TO_CATEGORY.items():
        # word-boundary-ish match, tolerant of . and + in skill names (c++, node.js)
        pattern = re.escape(skill)
        if re.search(rf"(?<![a-z0-9]){pattern}(?![a-z0-9])", text_lower):
            found.setdefault(category, [])
            display = skill.title() if skill.isalpha() else skill
            found[category].append(display)
    for cat in found:
        found[cat] = sorted(set(found[cat]))
    return found


def extract_education(section_text):
    if not section_text:
        return []
    entries = []
    blocks = re.split(r"\n{1,}", section_text)
    current = []

    def flush(block_lines):
        block = " ".join(block_lines).strip()
        if not block:
            return None
        degree = None
        for kw in DEGREE_KEYWORDS:
            if kw in block.lower():
                degree = kw.upper() if len(kw) <= 6 else kw.title()
                break
        years = YEAR_RANGE_RE.search(block)
        year_val = years.group(0) if years else None
        if not year_val:
            y = YEAR_RE.findall(block)
            year_val = y[-1] if y else None
        cgpa = CGPA_RE.search(block)
        pct = PERCENT_RE.search(block)
        return {
            "raw": block,
            "degree": degree,
            "year": year_val,
            "cgpa": cgpa.group(1) if cgpa else None,
            "percentage": pct.group(1) + "%" if pct else None,
        }

    for line in blocks:
        line = line.strip()
        if not line:
            if current:
                entry = flush(current)
                if entry:
                    entries.append(entry)
                current = []
            continue
        current.append(line)
    if current:
        entry = flush(current)
        if entry:
            entries.append(entry)

    return entries


def extract_experience(section_text):
    if not section_text:
        return []
    entries = []
    lines = [l for l in section_text.split("\n") if l.strip()]
    current_block = []

    def flush(block_lines):
        if not block_lines:
            return None
        block = "\n".join(block_lines)

        # Prefer a precise month-year range ("March 2024 - Present") over a
        # bare year range ("2024 - Present"), since it lets us compute an
        # accurate duration down to months/days.
        month_range = MONTH_YEAR_RANGE_RE.search(block)
        year_range = YEAR_RANGE_RE.search(block)

        duration_text = None
        duration_calculated = None

        if month_range:
            duration_text = month_range.group(0)
            duration_calculated = calculate_duration(month_range.group(1), month_range.group(2))
        elif year_range:
            duration_text = year_range.group(0)
            duration_calculated = calculate_duration(year_range.group(1), year_range.group(3))

        bullets = [l.strip("-•* \t") for l in block_lines if l.strip().startswith(("-", "•", "*"))]
        header_line = block_lines[0]
        return {
            "raw_header": header_line.strip(),
            "duration": duration_text,
            "duration_calculated": duration_calculated,
            "details": bullets if bullets else None,
            "full_text": block.strip(),
        }

    for line in lines:
        # New entry heuristic: line contains a date range or looks like "Title, Company"
        if (MONTH_YEAR_RANGE_RE.search(line) or YEAR_RANGE_RE.search(line)) and current_block:
            current_block.append(line)
            continue
        if current_block and not line.strip().startswith(("-", "•", "*")) and \
           (re.search(r"\b(20\d{2}|19\d{2})\b", line) or len(line.split()) <= 8) and \
           len(current_block) > 1:
            entry = flush(current_block)
            if entry:
                entries.append(entry)
            current_block = [line]
        else:
            current_block.append(line)

    if current_block:
        entry = flush(current_block)
        if entry:
            entries.append(entry)

    return entries


def calculate_total_experience(experience_entries):
    """Sum up total_days across all experience entries with a calculated
    duration, and express it as years/months. Overlapping roles will be
    double-counted (treated as separate durations), matching how most
    resume/ATS total-experience calculators behave."""
    total_days = 0
    for entry in experience_entries:
        calc = entry.get("duration_calculated")
        if calc:
            total_days += calc["total_days"]

    if total_days == 0:
        return None

    years = total_days // 365
    remaining_days = total_days % 365
    months = remaining_days // 30
    days = remaining_days % 30

    parts = []
    if years:
        parts.append(f"{years} year{'s' if years != 1 else ''}")
    if months:
        parts.append(f"{months} month{'s' if months != 1 else ''}")
    if not parts:
        parts.append(f"{days} day{'s' if days != 1 else ''}")

    return {
        "total_days": total_days,
        "years": int(years),
        "months": int(months),
        "human_readable": " ".join(parts),
    }


def extract_projects(section_text):
    if not section_text:
        return []
    blocks = re.split(r"\n(?=[A-Z][^\n]{0,80}$)", section_text, flags=re.M)
    projects = []
    for block in blocks:
        block = block.strip()
        if not block:
            continue
        lines = block.split("\n")
        title = lines[0].strip("-•* \t")
        bullets = [l.strip("-•* \t") for l in lines[1:] if l.strip()]
        projects.append({"title": title, "details": bullets})
    return projects


def extract_certifications(section_text):
    if not section_text:
        return []
    lines = [l.strip("-•* \t") for l in section_text.split("\n") if l.strip()]
    return lines


def extract_summary(section_text):
    return section_text.strip() if section_text else None


# --------------------------------------------------------------------------
# Main parse entrypoint
# --------------------------------------------------------------------------

def parse_resume(path: str) -> dict:
    raw_text = extract_text(path)
    if not raw_text or not raw_text.strip():
        raise ValueError("No extractable text found in this file. "
                          "It may be a scanned image — try an OCR'd version.")

    sections, lines = segment_sections(raw_text)
    experience_entries = extract_experience(sections.get("experience"))

    result = {
        "name": extract_name(lines),
        "contact_info": extract_contact_info(raw_text),
        "summary": extract_summary(sections.get("summary")),
        "education": extract_education(sections.get("education")),
        "experience": experience_entries,
        "total_experience": calculate_total_experience(experience_entries),
        "projects": extract_projects(sections.get("projects")),
        "skills": extract_skills(raw_text),
        "certifications": extract_certifications(sections.get("certifications")),
        "achievements": extract_certifications(sections.get("achievements")),
        "sections_detected": list(sections.keys()),
        "raw_text_length": len(raw_text),
    }
    return result


if __name__ == "__main__":
    import sys
    if len(sys.argv) < 2:
        print("Usage: python parser.py <path_to_resume>")
        sys.exit(1)
    data = parse_resume(sys.argv[1])
    print(json.dumps(data, indent=2))